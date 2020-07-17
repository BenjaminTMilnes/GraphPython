from graphe.core import *
from docx import Document
from docx.shared import Pt, Mm, Cm, Inches, RGBColor
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import logging

logging.basicConfig(level=logging.DEBUG)

alignments = {
    "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left-justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "right-justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "centre-justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "center-justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "centred": WD_ALIGN_PARAGRAPH.CENTER,
    "centered": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


class WordExportContext(object):
    def __init__(self, dx):
        self.logger = logging.getLogger(__name__)

        self.dx = dx

        self.n = 0
        self.currentSection = self.dx.sections[0]
        self.currentParagraph = None
        self.lastParagraph = None
        self.currentRun = None

    def addSection(self, pageWidth, pageHeight, topMargin, rightMargin, bottomMargin, leftMargin):
        self.logger.debug("Adding section to document.")

        if self.n > 0:
            self.currentSection = self.dx.add_section(WD_SECTION.NEW_PAGE)

        self.n = 1

        self.currentSection.page_width = pageWidth
        self.currentSection.page_height = pageHeight
        self.currentSection.top_margin = topMargin
        self.currentSection.right_margin = rightMargin
        self.currentSection.bottom_margin = bottomMargin
        self.currentSection.left_margin = leftMargin

    def addHeading(self, level, textAlignment, marginTop, marginBottom):
        self.logger.debug("Adding heading to document.")

        self.currentParagraph = self.dx.add_paragraph("")
        self.currentParagraph.alignment = alignments.get(textAlignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.currentParagraph.paragraph_format.space_before = marginTop
        self.currentParagraph.paragraph_format.space_after = marginBottom

        return

        if level > 9:
            level = 9
        if level < 1:
            level = 1
        self.currentParagraph = self.dx.add_heading("", level)

    def addParagraph(self, textAlignment, marginTop, marginBottom, lineHeight, textIndentation):
        self.logger.debug("Adding paragraph to document.")

        self.currentParagraph = self.dx.add_paragraph("")
        self.currentParagraph.alignment = alignments.get(textAlignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.currentParagraph.paragraph_format.space_before = marginTop
        self.currentParagraph.paragraph_format.space_after = marginBottom
        self.currentParagraph.paragraph_format.line_spacing = lineHeight
        self.currentParagraph.paragraph_format.first_line_indent = textIndentation

    def addRun(self, text, fontName, fontHeight, bold=False, italic=False, underline=False, strikethrough=False, fontVariant="none"):
        self.logger.debug("Adding text \"{}\" to document.".format(text))

        self.currentRun = self.currentParagraph.add_run(text)
        self.currentRun.font.name = fontName
        self.currentRun.font.size = fontHeight
        self.currentRun.font.bold = bold
        self.currentRun.font.italic = italic
        self.currentRun.font.underline = underline
        self.currentRun.font.strike = strikethrough
        self.currentRun.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        self.currentRun.font.small_caps = True if fontVariant in ["small-capitals", "small-caps"] else False

    def addLineBreak(self):
        self.logger.debug("Adding line break to document.")
        self.currentRun.add_break()

    def enterSectionHeader(self):
        self.lastParagraph = self.currentParagraph

        header = self.currentSection.header
        self.currentParagraph = header.paragraphs[0]
        self.currentParagraph.text = "lkj"

    def exitSectionHeader(self):
        self.currentParagraph = self.lastParagraph


class WordExporter(object):
    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def _getLength(self, length):
        self.logger.debug(length)

        if length.unit == "mm":
            return Mm(length.number)
        if length.unit == "cm":
            return Cm(length.number)
        if length.unit == "dm":
            return Cm(length.number * 10)
        if length.unit == "m":
            return Cm(length.number * 100)
        if length.unit == "pt":
            return Pt(length.number)
        if length.unit == "in":
            return Inches(length.number)

    def exportDocument(self, document, filePath):

        dx = Document()
        context = WordExportContext(dx)

        for section in document.sections:
            self.exportSection(section, document, context)

        dx.save(filePath)

    def exportSection(self, section, document, context):
        pageWidth = self._getLength(section.styleProperties.get("page-width", GLength(12.85, "cm")))
        pageHeight = self._getLength(section.styleProperties.get("page-height", GLength(19.84, "cm")))
        marginTop = self._getLength(section.styleProperties.get("margin-top", GLength(1.5, "cm")))
        marginRight = self._getLength(section.styleProperties.get("margin-right", GLength(1.5, "cm")))
        marginBottom = self._getLength(section.styleProperties.get("margin-bottom", GLength(1.5, "cm")))
        marginLeft = self._getLength(section.styleProperties.get("margin-left", GLength(1.5, "cm")))

        context.addSection(pageWidth, pageHeight, marginTop, marginRight, marginBottom, marginLeft)

        self.logger.info("Page template reference: '{0}'".format(section.pageTemplateReference))

        if section.pageTemplate != None and section.pageTemplate.header != None:
            context.enterSectionHeader()
            self.exportPageElements(section.pageTemplate.header.subelements, document, context)
            context.exitSectionHeader()

        self.exportPageElements(section.subelements, document, context)

    def exportPageElements(self, pageElements, document, context):
        for pageElement in pageElements:
            self.exportPageElement(pageElement, document, context)

    def exportPageElement(self, pageElement, document, context):

        if isinstance(pageElement, GParagraph):
            textAlignment = pageElement.styleProperties.get("text-alignment", "left")
            marginTop = self._getLength(pageElement.styleProperties.get("margin-top", GLength(0, "pt")))
            marginBottom = self._getLength(pageElement.styleProperties.get("margin-bottom", GLength(0, "pt")))
            textIndentation = self._getLength(pageElement.styleProperties.get("text-indentation", GLength(0, "pt")))
            lineHeight = self._getLength(pageElement.styleProperties.get("line-height", GLength(12, "pt")))

            context.addParagraph(textAlignment, marginTop, marginBottom, lineHeight, textIndentation)
            self.exportPageElements(pageElement.subelements, document, context)

        if isinstance(pageElement, GHeading):
            textAlignment = pageElement.styleProperties.get("text-alignment", "left")
            marginTop = self._getLength(pageElement.styleProperties.get("margin-top", GLength(0, "pt")))
            marginBottom = self._getLength(pageElement.styleProperties.get("margin-bottom", GLength(0, "pt")))

            context.addHeading(pageElement.level, textAlignment, marginTop, marginBottom)
            self.exportPageElements(pageElement.subelements, document, context)

        if isinstance(pageElement, GDivision) or isinstance(pageElement, GDefinitionListTerm) or isinstance(pageElement, GDefinitionListDefinition):
            textAlignment = pageElement.styleProperties.get("text-alignment", "left")
            marginTop = self._getLength(pageElement.styleProperties.get("margin-top", GLength(0, "pt")))
            marginBottom = self._getLength(pageElement.styleProperties.get("margin-bottom", GLength(0, "pt")))
            textIndentation = self._getLength(pageElement.styleProperties.get("text-indentation", GLength(0, "pt")))
            lineHeight = self._getLength(pageElement.styleProperties.get("line-height", GLength(12, "pt")))

            context.addParagraph(textAlignment, marginTop, marginBottom, lineHeight, textIndentation)
            self.exportPageElements(pageElement.subelements, document, context)
            
        if isinstance(pageElement, GDefinitionList):
            self.exportPageElements(pageElement.subelements, document, context)

        if isinstance(pageElement, GVariable):
            if pageElement.name == "title":
                context.addRun(document.title)
            if pageElement.name == "subtitle":
                context.addRun(document.subtitle)
            if pageElement.name == "authorName":
                context.addRun(document.authorName)
            if pageElement.name == "currentYear":
                context.addRun(str(datetime.now().year))

        if isinstance(pageElement, GHyperlink):
            self.exportPageElements(pageElement.subelements, document, context)

        if isinstance(pageElement, GItalic):
            self.exportPageElements(pageElement.subelements, document, context)
            
        if isinstance(pageElement, GBold):
            self.exportPageElements(pageElement.subelements, document, context)

        if isinstance(pageElement, GTextElement):
            fontName = pageElement.styleProperties.get("font-name", "Times New Roman")
            fontVariant = pageElement.styleProperties.get("font-variant", "none")
            fontHeight = self._getLength(pageElement.styleProperties.get("font-height", GLength(12, "pt")))
            fontSlant = True if pageElement.styleProperties.get("font-slant", "none") == "italic" else False
            fontWeight = True if pageElement.styleProperties.get("font-weight", "none") == "bold" else False

            context.addRun(pageElement.text, fontName, fontHeight, fontWeight, fontSlant, False, False, fontVariant)

        if isinstance(pageElement, GLineBreak):
            context.addLineBreak()
